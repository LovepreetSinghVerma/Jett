from selenium import webdriver
from selenium.webdriver.common.by import By
import time
from docx import Document
from selenium.webdriver.common.keys import Keys
import re

def is_heading(paragraph):
    # Check if the paragraph is a heading or subheading
    style = paragraph.style.name.lower()
    return style.startswith("heading")

def select_300_words(paragraphs, index):
    word_count = 0
    selected_text = []

    while index < len(paragraphs):
        paragraph = paragraphs[index]

        if is_heading(paragraph):
            index += 1
            continue

        for run in paragraph.runs:
            words = re.findall(r'\b\w+\b', run.text)
            for word in words:
                selected_text.append(word)
                word_count += 1
                if word_count >= 300:
                    # Find the next full stop
                    rest_of_paragraph = run.text.split(word)[-1]
                    next_full_stop_index = rest_of_paragraph.find('.')
                    if next_full_stop_index != -1:
                        selected_text.append(rest_of_paragraph[:next_full_stop_index + 1])
                    return ' '.join(selected_text)

        index += 1

    return ' '.join(selected_text)


driver = webdriver.Chrome() 
driver.get('https://stealthwriter.ai/')
time.sleep(2)
# Sign in
driver.find_element(By.XPATH, '/html/body/div[1]/div/nav/div[2]/div/div[2]/a').click()
time.sleep(2)

email = driver.find_element(By.XPATH, '/html/body/div[1]/div/main/main/div/div[1]/form/div[1]/div[1]/div/input').send_keys('parasverma72683@gmail.com')
password = driver.find_element(By.XPATH, '/html/body/div[1]/div/main/main/div/div[1]/form/div[1]/div[2]/div/input').send_keys('DualNature')
time.sleep(20)
# Login
login = driver.find_element(By.XPATH,'/html/body/div[1]/div/main/main/div/div[1]/form/button').click()
print(login)
time.sleep(5)
# Feed data
def humanizer(data):
    box = driver.find_element(By.XPATH,'/html/body/div[1]/div/main/div[2]/div/main/div[1]/form/div[2]/div/div/textarea[1]').click()
    time.sleep(2)
    webdriver.ActionChains(driver).key_down(Keys.CONTROL).send_keys('a').key_up(Keys.CONTROL).perform()
    time.sleep(1)
    webdriver.ActionChains(driver).send_keys(Keys.BACKSPACE).perform()
    time.sleep(1)
    driver.find_element(By.XPATH,'/html/body/div[1]/div/main/div[2]/div/main/div[1]/form/div[2]/div/div/textarea[1]').send_keys(f'{data}') 
    
    # Humanize
    driver.find_element(By.XPATH,'/html/body/div[1]/div/main/div[2]/div/main/div[1]/form/div[3]/div/button').click()
    time.sleep(4)

    # Copy rephrased text
    rephrased_text = driver.find_element(By.XPATH,'/html/body/div[1]/div/main/div[2]/div/main/div[2]/div/div[3]/div/div/div[1]').text
    return rephrased_text

def replace_text_in_document(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)

def main(file_path):
    doc = Document(file_path)
    paragraphs = doc.paragraphs
    index = 0    
    while index < len(paragraphs):
        selected_text = select_300_words(paragraphs, index)
        if not selected_text:
            break
        rephrased_text = humanizer(selected_text)
        replace_text_in_document(doc, selected_text, rephrased_text)
        index += 1

    doc.save(r"C:\Users\HP\Desktop\Rephrased_Document.docx")
    driver.quit()

file_path = r"C:\Users\HP\Desktop\Testing.docx"
main(file_path)
